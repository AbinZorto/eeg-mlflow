#!/usr/bin/env python3
"""
pandoc_prep.py — Prepare a LaTeX manuscript for clean pandoc → Word conversion,
and format the resulting .docx to match journal submission guidelines.

== LATeX PREP MODE (default) ==
Fixes all known pandoc conversion issues:
  1. Converts longtable/xltabular → table+tabular (pandoc can't handle them)
  2. Replaces \\ref{label} with literal numbers (auto-detected from doc order)
  3. Adds "Figure N:" / "Table N:" prefixes to captions
  4. Strips >{\\raggedright}/>{\\raggedleft} column spec modifiers
  5. Converts inline math in table cells to unicode/plain text
  6. Removes \\addlinespace, \\arraybackslash, \\resizebox wrappers
  7. Transforms elsarticle preamble → article + geometry (optional)
  8. Removes elsarticle-specific commands (\\credit, \\cormark, \\ead, etc.)

== DOCX FORMAT MODE (--format-docx) ==
Post-processes a .docx to match journal guidelines:
  - Page size, margins, orientation
  - Font family / size (body, headings, captions, tables)
  - Line spacing (single, 1.5, double)
  - Paragraph alignment, indentation, spacing
  - Line numbering (continuous or per-page)
  - Page numbers (header/footer, alignment)
  - Caption styling
  - References hanging indent
  - Table font sizing
  - Widow/orphan control
  - Custom header/footer text
  - Title-page (different first page) support

Usage (LaTeX prep):
    python3 pandoc_prep.py input.tex -o output.tex --elsarticle

Usage (DOCX format):
    python3 pandoc_prep.py --format-docx input.docx --journal eswa
    python3 pandoc_prep.py --format-docx input.docx \\
        --font "Times New Roman" --font-size 12 --line-spacing 2.0 \\
        --margins 1 --line-numbers --page-numbers
"""

import re
import argparse
import sys
import os
from collections import OrderedDict


def resolve_inputs(content, base_dir, input_paths=None, max_depth=3):
    """Resolve \\input{} and \\include{} commands in content, expanding included files."""
    if max_depth <= 0:
        return content

    if input_paths is None:
        input_paths = [base_dir]

    def replace_input(match):
        args = match.group(1).strip()
        fname = args.split('}')[0] if '}' in args else args
        fname = fname.strip('{}')

        for path_dir in input_paths:
            candidate = os.path.join(path_dir, fname)
            if not candidate.endswith('.tex'):
                candidate += '.tex'
            if os.path.exists(candidate):
                try:
                    with open(candidate, 'r') as f:
                        sub_content = f.read()
                    sub_content = resolve_inputs(sub_content, base_dir, input_paths, max_depth - 1)
                    return sub_content
                except Exception:
                    pass
        return match.group(0)

    content = re.sub(r'\\(?:input|include)\{([^}]+(?:\{[^}]*\}[^}]*)*)\}',
                     replace_input, content)
    return content

def build_input_paths(content, base_dir):
    """Detect \\graphicspath and \\input@path from preamble to build search paths."""
    paths = [base_dir]
    for m in re.finditer(r'\\graphicspath\{\{(.+?)\}\}', content):
        for p in m.group(1).split(','):
            p = p.strip().strip('"').strip("'")
            resolved = os.path.join(base_dir, p)
            if os.path.isdir(resolved):
                paths.append(resolved)
    for m in re.finditer(r'\\def\\input@path\{\{(.+?)\}\}', content):
        for p in m.group(1).split(','):
            p = p.strip().strip('"').strip("'")
            resolved = os.path.join(base_dir, p)
            if os.path.isdir(resolved):
                paths.append(resolved)
    # Deduplicate while preserving order
    seen = set()
    unique = []
    for p in paths:
        if p not in seen:
            seen.add(p)
            unique.append(p)
    return unique


def detect_labels(content):
    """Scan document for \\label{} commands, return ordered list of (label, type)."""
    labels = []
    skip_prefixes = ('sec:', 'subsec:', 'eq:', 'alg:', 'lst:', 'app:', 'fn:')
    
    for m in re.finditer(r'\\label\{([^}]+)\}', content):
        label = m.group(1)
        
        # Skip section/equation/list labels
        if label.startswith(skip_prefixes):
            continue
        
        # Look backwards for context: \caption inside figure/table/longtable/xltabular
        before = content[max(0, m.start()-800):m.start()]
        if re.search(r'\\begin\{figure\}', before, re.DOTALL):
            labels.append((label, 'figure'))
        elif re.search(r'\\begin\{table\}', before, re.DOTALL):
            labels.append((label, 'table'))
        elif re.search(r'\\begin\{(?:longtable|xltabular)\}', before, re.DOTALL):
            labels.append((label, 'table'))
        elif 'fig:' in label:
            labels.append((label, 'figure'))
        elif 'tab:' in label:
            labels.append((label, 'table'))
        else:
            # Check contextual proximity: is there a \caption within 300 chars?
            before_short = content[max(0, m.start()-300):m.start()]
            if re.search(r'\\caption\{', before_short, re.DOTALL):
                if re.search(r'\\includegraphics', before_short, re.DOTALL):
                    labels.append((label, 'figure'))
                elif re.search(r'\\begin\{tabular\}', before_short, re.DOTALL):
                    labels.append((label, 'table'))
                else:
                    # Can't determine — skip
                    pass
            else:
                # No nearby caption — likely not a figure/table label
                pass
    return labels


def build_label_map(labels):
    """Assign sequential numbers: Figure 1,2,3... Table 1,2,3..."""
    fig_map = OrderedDict()
    tab_map = OrderedDict()
    fig_count = 0
    tab_count = 0
    for label, kind in labels:
        if kind == 'figure':
            fig_count += 1
            fig_map[label] = fig_count
        elif kind == 'table':
            tab_count += 1
            tab_map[label] = tab_count
    return fig_map, tab_map


def fix_refs(content, fig_map, tab_map):
    """Replace \\ref{label} with literal numbers."""
    all_labels = {**fig_map, **tab_map}

    def replace_ref(match):
        label = match.group(1)
        if label in all_labels:
            return str(all_labels[label])
        # Label not found — leave as-is with a warning comment
        return f"REF:{label}"

    content = re.sub(r'\\ref\{([^}]+)\}', replace_ref, content)
    return content


def fix_captions(content, fig_map, tab_map):
    """Add 'Figure N: ' or 'Table N: ' prefixes to \\caption{} commands."""
    lines = content.split('\n')
    new_lines = []

    for i, line in enumerate(lines):
        cap_match = re.match(r'(\s*\\caption\{)(.+)', line)
        if cap_match:
            cap_start = cap_match.group(1)
            cap_body = cap_match.group(2)

            # Find label in next few lines
            label_found = None
            for j in range(i, min(i + 5, len(lines))):
                l_match = re.search(r'\\label\{([^}]+)\}', lines[j])
                if l_match:
                    label_found = l_match.group(1)
                    break

            if label_found and label_found in fig_map:
                prefix = f"Figure {fig_map[label_found]}: "
            elif label_found and label_found in tab_map:
                prefix = f"Table {tab_map[label_found]}: "
            else:
                prefix = None

            if prefix and not re.match(r'^(Figure|Table)\s+\d+[:. ]', cap_body.strip()):
                line = cap_start + prefix + cap_body

        new_lines.append(line)
    return '\n'.join(new_lines)


def fix_table_columns(content):
    """Remove >{\\raggedright}, >{\\raggedleft}, \\arraybackslash from column specs."""
    # Handle both standard and raw-LaTeX representations
    content = content.replace('>{\\raggedright}', '')
    content = content.replace('>{\\raggedleft}', '')
    content = content.replace('>{\\raggedright\\arraybackslash}', '')
    content = content.replace('>{\\raggedleft\\arraybackslash}', '')
    content = content.replace('\\arraybackslash', '')
    # Also clean up any leftover empty >{} modifiers
    content = re.sub(r'>\{\}', '', content)
    return content


def fix_table_math(content):
    """Convert inline math in table cells to unicode/plain text equivalents."""
    # Only fix inside tabular environments (body-text math is fine)
    table_pattern = re.compile(r'(\\begin\{(?:tabular|tabularx|longtable|xltabular)\}.*?\\end\{(?:tabular|tabularx|longtable|xltabular)\})', re.DOTALL)

    math_replacements = [
        (r'\$\\?pm\$', '\u00b1'),           # ±
        (r'\$\\mp\$', '\u2213'),             # ∓
        (r'\$\\approx\$', '\u2248'),         # ≈
        (r'\$\\sim\$', '~'),
        (r'\$\\ge\$', '\u2265'),             # ≥
        (r'\$\\le\$', '\u2264'),             # ≤
        (r'\$\\times\$', '\u00d7'),          # ×
        (r'\$\\alpha\$', '\u03b1'),          # α
        (r'\$\\beta\$', '\u03b2'),           # β
        (r'\$\\gamma\$', '\u03b3'),          # γ
        (r'\$\\delta\$', '\u03b4'),          # δ
        (r'\$\\Delta\$', '\u0394'),          # Δ
        (r'\$\\mu\$', '\u03bc'),             # μ
        (r'\$\\sigma\$', '\u03c3'),          # σ
        (r'\$\\dagger\$', '\u2020'),         # †
        (r'\$\\text\{([^}]*)\}\$', r'\1'),  # \text{...} → plain text
        (r'\$\\mathcal\{N\}\(0,\s*1\)\$', 'N(0,1)'),
        (r'\$\\mathcal\{N\}\$', 'N'),
        (r'\$\\mathrm\{([^}]*)\}\$', r'\1'),
        # Single-letter math: $X$ where X is a single char
        (r'\$([fFkpPnNdDrR])\$', r'\1'),
    ]

    def fix_math_in_block(block):
        for pattern, replacement in math_replacements:
            block = re.sub(pattern, replacement, block)
        return block

    matches = list(table_pattern.finditer(content))
    for m in reversed(matches):
        fixed = fix_math_in_block(m.group(0))
        content = content[:m.start()] + fixed + content[m.end():]

    return content


def fix_addlinespace(content):
    """Remove \\addlinespace commands (pandoc can't handle them)."""
    return re.sub(r'\\addlinespace(\[[^\]]*\])?', '', content)


def fix_resizebox(content):
    """Remove \\resizebox{...}{...}{ from around tabular environments.
    Also removes the matching closing brace.
    pandoc can't handle resizebox-wrapped tabular content."""
    # Remove \resizebox{width}{height}{ 
    content = re.sub(r'\\resizebox\{[^}]*\}\{[^}]*\}\{', '', content)
    # Remove the matching closing } — it appears right after \end{tabular}
    content = re.sub(r'(\\end\{tabular\})\s*\}', r'\1', content)
    return content


def fix_longtables(content):
    """Convert longtable/xltabular environments to simple table+tabular."""
    import re as _re

    def _balanced_braces(text, start_pos):
        """Extract brace-balanced content starting at '{' at start_pos."""
        if start_pos >= len(text) or text[start_pos] != '{':
            return '', start_pos
        depth = 1
        i = start_pos + 1
        while i < len(text) and depth > 0:
            if text[i] == '{':
                depth += 1
            elif text[i] == '}':
                depth -= 1
            i += 1
        return text[start_pos+1:i-1], i

    def _extract_caption_label(inner):
        """Extract caption text, label, AND header rows from xltabular preamble.
        
        Returns (caption, label, header_rows) where header_rows is everything
        between the caption/label end and \\endfirsthead (the \\toprule, column
        headings, and \\midrule that form the actual table header).
        """
        caption = ''
        label = ''
        header_rows = ''
        efh_idx = inner.find('\\endfirsthead')
        if efh_idx >= 0:
            preamble = inner[:efh_idx]
        else:
            preamble = inner
        
        # Find caption
        cap_start = preamble.find('\\caption{')
        if cap_start >= 0:
            cap_body, cap_end = _balanced_braces(preamble, cap_start + len('\\caption'))
            caption = cap_body.strip()
            remaining = preamble[cap_end:].strip()
        else:
            remaining = preamble
        
        # Find label in remaining (before or after caption)
        lab_match = _re.search(r'\\label\{([^}]+)\}', remaining)
        if lab_match:
            label = lab_match.group(1)
            remaining = remaining[lab_match.end():].strip()
        
        # Everything left is the header rows (\toprule, column headings, \midrule)
        # Clean up: strip trailing/leading \\ (line terminators around markers)
        header_rows = remaining.strip()
        if header_rows.endswith('\\\\'):
            header_rows = header_rows[:-2].strip()
        if header_rows.startswith('\\\\'):
            header_rows = header_rows[2:].strip()
        
        return caption, label, header_rows

    def _extract_body(inner):
        for marker in ['\\endlastfoot', '\\endfoot']:
            idx = inner.find(marker)
            if idx >= 0:
                return inner[idx + len(marker):].strip()
        return inner.strip()

    def _convert_x_columns(colspec):
        if 'X' not in colspec:
            return colspec
        x_count = len(_re.findall(r'\bX\b', colspec))
        if x_count == 0:
            return colspec
        non_x = sum(float(m.group(1)) for m in _re.finditer(r'p\{([0-9.]+)\\textwidth\}', colspec))
        remaining = max(0.05, 0.95 - non_x)
        width = remaining / x_count
        # Use lambda so re.sub doesn't interpret \t in the replacement as a tab
        return _re.sub(r'\bX\b',
                       lambda m: f'p{{{width:.2f}\\textwidth}}',
                       colspec)

    def _extract_env(env_name, content):
        pattern = _re.compile(r'\\begin\{' + env_name + r'\}', _re.DOTALL)
        out_parts = []
        last_end = 0
        for m in pattern.finditer(content):
            out_parts.append(content[last_end:m.start()])
            pos = m.end()
            specs = []
            for _ in range(2 if env_name == 'xltabular' else 1):
                while pos < len(content) and content[pos] in ' \t\n\r':
                    pos += 1
                if pos < len(content) and content[pos] == '{':
                    spec, pos = _balanced_braces(content, pos)
                    specs.append(spec)
                else:
                    break
            end_marker = '\\end{' + env_name + '}'
            end_pos = content.find(end_marker, pos)
            if end_pos < 0:
                end_pos = pos
            inner = content[pos:end_pos]
            colspec = ''
            if env_name == 'xltabular' and len(specs) >= 2:
                colspec = _convert_x_columns(specs[1])
            elif specs:
                colspec = specs[0]
            caption, label, header_rows = _extract_caption_label(inner)
            body = _extract_body(inner)
            result = '\\begin{table}[htbp]\n\\centering\n'
            if caption:
                result += f'\\caption{{{caption}}}\n'
            if label:
                result += f'\\label{{{label}}}\n'
            result += f'\\begin{{tabular}}{{{colspec}}}\n'
            if header_rows:
                result += header_rows + '\n'
            result += body
            # Ensure bottom rule at end of body for proper booktabs
            body_end = body.rstrip()
            if not body_end.endswith('\\bottomrule'):
                if body_end.endswith('\\\\'):
                    result = result.rstrip() + '\n\\bottomrule'
                else:
                    result = result.rstrip() + '\\tabularnewline\n\\bottomrule'
            result += '\n\\end{tabular}\n\\end{table}'
            out_parts.append(result)
            last_end = end_pos + len(end_marker)
        out_parts.append(content[last_end:])
        return ''.join(out_parts)

    content = _extract_env('longtable', content)
    content = _extract_env('xltabular', content)
    return content


def fix_double_prefixes(content):
    """Fix doubled prefixes like 'Figure Figure 1' → 'Figure 1'."""
    for prefix in ['Figure', 'Figures', 'Table', 'Tables']:
        pattern = re.compile(rf'{prefix}\s*{prefix}\s+(\d+)')
        content = re.sub(pattern, rf'{prefix} \1', content)
    return content


def transform_elsarticle_preamble(content):
    """Convert elsarticle preamble to article class + geometry."""
    # Replace documentclass
    content = re.sub(
        r'\\documentclass(\[[^\]]*\])?\{elsarticle\}',
        r'\\documentclass[12pt]{article}\n\\usepackage[utf8]{inputenc}\n\\usepackage[T1]{fontenc}\n\\usepackage[margin=1in]{geometry}\n\\usepackage{setspace}\\setstretch{1.5}',
        content
    )

    # Remove elsarticle-specific commands
    removals = [
        r'\\journal\{[^}]*\}',
        r'\\begin\{frontmatter\}',
        r'\\end\{frontmatter\}',
        r'\\cormark\[[^\]]*\]',
        r'\\ead\{[^}]*\}',
        r'\\credit\{[^}]*\}',
        r'\\cortext\[[^\]]*\].*?(?=\n)',  # remove cortext lines
        r'\\printcredits',
        r'\\MSC\[[^\]]*\].*?(?=\n)',
    ]
    for pattern in removals:
        content = re.sub(pattern, '', content)

    # \affiliation is multi-line with nested braces
    # Match from \affiliation through to end of command (braces may be nested)
    content = re.sub(
        r'\\affiliation\[[^\]]*\].*?(?=\n\s*\n|\\[a-z])',
        '',
        content,
        flags=re.DOTALL
    )
    content = re.sub(
        r'\\affiliation\{.*?(?=\n\s*\n|\\[a-z])',
        '',
        content,
        flags=re.DOTALL
    )

    # Simplify \author[inst1]{Name} → \author{Name}
    content = re.sub(r'\\author\[[^\]]*\]\{', r'\\author{', content)

    # Simplify custom environments for pandoc compatibility
    content = re.sub(
        r'\\newenvironment\{paperwidefigure\}\{(.+?)\}\{(.+?)\}',
        r'\\newenvironment{paperwidefigure}{\\begin{figure}[htbp]}{\\end{figure}}',
        content
    )
    content = re.sub(
        r'\\newenvironment\{paperinlinefigure\}\{(.+?)\}\{(.+?)\}',
        r'\\newenvironment{paperinlinefigure}{\\begin{figure}[htbp]}{\\end{figure}}',
        content
    )
    content = re.sub(
        r'\\newenvironment\{paperwidetable\}\{(.+?)\}\{(.+?)\}',
        r'\\newenvironment{paperwidetable}{\\begin{table}[htbp]}{\\end{table}}',
        content
    )

    return content


def extract_label_map_text(labels, fig_map, tab_map):
    """Generate readable label → number mapping for user reference."""
    lines = []
    lines.append("LABEL → NUMBER MAPPING (auto-detected):")
    for label, num in fig_map.items():
        lines.append(f"  Figure {num}: {label}")
    for label, num in tab_map.items():
        lines.append(f"  Table {num}: {label}")
    return '\n'.join(lines)


# =============================================================================
# DOCX FORMATTING MODULE
# =============================================================================

JOURNAL_PRESETS = {
    'eswa': {
        'description': 'Expert Systems with Applications (Elsevier)',
        'page_size': 'letter',
        'font': 'Times New Roman',
        'font_size': 12,
        'line_spacing': 2.0,
        'margin_top': 1.0,
        'margin_bottom': 1.0,
        'margin_left': 1.0,
        'margin_right': 1.0,
        'line_numbers': True,
        'line_numbers_restart': 'page',
        'page_numbers': True,
        'page_numbers_position': 'footer',
        'page_numbers_align': 'center',
        'justify_body': True,
        'space_after': 0,
        'space_before': 0,
        'first_line_indent': 0,
        'caption_font_size': 10,
        'caption_align': 'center',
        'references_hanging_indent': 0.5,
        'table_font_size': 9,
        'table_borders': 'booktabs',
        'heading_bold': True,
        'heading_size': None,
        'widow_orphan': True,
    },
    'dss': {
        'description': 'Decision Support Systems (Elsevier, Vancouver citations)',
        'page_size': 'letter',
        'font': 'Times New Roman',
        'font_size': 12,
        'line_spacing': 2.0,
        'margin_top': 1.0,
        'margin_bottom': 1.0,
        'margin_left': 1.0,
        'margin_right': 1.0,
        'line_numbers': True,
        'line_numbers_restart': 'page',
        'page_numbers': True,
        'page_numbers_position': 'footer',
        'page_numbers_align': 'center',
        'justify_body': True,
        'space_after': 0,
        'space_before': 0,
        'first_line_indent': 0,
        'caption_font_size': 10,
        'caption_align': 'center',
        'references_hanging_indent': 0.5,
        'table_font_size': 9,
        'table_borders': 'booktabs',
        'heading_bold': True,
        'heading_size': None,
        'widow_orphan': True,
    },
    'jtsa': {
        'description': 'Journal of Time Series Analysis',
        'page_size': 'letter',
        'font': 'Times New Roman',
        'font_size': 12,
        'line_spacing': 2.0,
        'margin_top': 1.0,
        'margin_bottom': 1.0,
        'margin_left': 1.0,
        'margin_right': 1.0,
        'line_numbers': True,
        'line_numbers_restart': 'page',
        'page_numbers': True,
        'page_numbers_position': 'footer',
        'page_numbers_align': 'right',
        'justify_body': True,
        'space_after': 0,
        'space_before': 0,
        'first_line_indent': 0.5,
        'caption_font_size': 11,
        'caption_align': 'center',
        'references_hanging_indent': 0.5,
        'table_font_size': 10,
        'table_borders': 'booktabs',
        'heading_bold': True,
        'heading_size': None,
        'widow_orphan': True,
    },
    'manuscript': {
        'description': 'Standard manuscript (double-spaced, TNR 12pt, 1" margins)',
        'page_size': 'letter',
        'font': 'Times New Roman',
        'font_size': 12,
        'line_spacing': 2.0,
        'margin_top': 1.0,
        'margin_bottom': 1.0,
        'margin_left': 1.0,
        'margin_right': 1.0,
        'line_numbers': False,
        'line_numbers_restart': 'page',
        'page_numbers': True,
        'page_numbers_position': 'footer',
        'page_numbers_align': 'center',
        'justify_body': False,
        'space_after': 0,
        'space_before': 0,
        'first_line_indent': 0.5,
        'caption_font_size': 11,
        'caption_align': 'center',
        'references_hanging_indent': 0.5,
        'table_font_size': 10,
        'table_borders': 'booktabs',
        'heading_bold': True,
        'heading_size': None,
        'widow_orphan': True,
    },
    'draft': {
        'description': 'Readable draft (1.15 spacing, 11pt, compact margins)',
        'page_size': 'letter',
        'font': 'Calibri',
        'font_size': 11,
        'line_spacing': 1.15,
        'margin_top': 0.8,
        'margin_bottom': 0.8,
        'margin_left': 1.0,
        'margin_right': 1.0,
        'line_numbers': False,
        'line_numbers_restart': 'continuous',
        'page_numbers': True,
        'page_numbers_position': 'footer',
        'page_numbers_align': 'center',
        'justify_body': True,
        'space_after': 8,
        'space_before': 0,
        'first_line_indent': 0,
        'caption_font_size': 10,
        'caption_align': 'center',
        'references_hanging_indent': 0.3,
        'table_font_size': 9,
        'table_borders': 'grid',
        'heading_bold': True,
        'heading_size': None,
        'widow_orphan': True,
    },
    'apa7': {
        'description': 'APA 7th Edition (double-spaced, TNR 12pt, 1" margins)',
        'page_size': 'letter',
        'font': 'Times New Roman',
        'font_size': 12,
        'line_spacing': 2.0,
        'margin_top': 1.0,
        'margin_bottom': 1.0,
        'margin_left': 1.0,
        'margin_right': 1.0,
        'line_numbers': False,
        'line_numbers_restart': 'page',
        'page_numbers': True,
        'page_numbers_position': 'header',
        'page_numbers_align': 'right',
        'justify_body': True,
        'space_after': 0,
        'space_before': 0,
        'first_line_indent': 0.5,
        'caption_font_size': 11,
        'caption_align': 'left',
        'references_hanging_indent': 0.5,
        'table_font_size': 10,
        'table_borders': 'booktabs',
        'heading_bold': True,
        'heading_size': None,
        'widow_orphan': True,
    },
}


PAGE_SIZES = {
    'letter':  (8.5, 11.0),   # width, height in inches
    'a4':      (8.27, 11.69),
    'legal':   (8.5, 14.0),
    'a5':      (5.83, 8.27),
    'tabloid': (11.0, 17.0),
}


def format_docx(docx_path, output_path, settings):
    """Apply journal formatting to a .docx file."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(docx_path)

    # Style names that pandoc uses for body text
    body_styles = ['Normal', 'Body Text', 'First Paragraph', 'Author', 'Abstract']
    caption_styles = ['Caption', 'Image Caption', 'Table Caption']

    # --- Page setup ---
    page_w, page_h = PAGE_SIZES.get(settings.get('page_size', 'letter'),
                                    PAGE_SIZES['letter'])
    for section in doc.sections:
        section.page_width = Inches(page_w)
        section.page_height = Inches(page_h)
        section.top_margin = Inches(settings.get('margin_top', 1.0))
        section.bottom_margin = Inches(settings.get('margin_bottom', 1.0))
        section.left_margin = Inches(settings.get('margin_left', 1.0))
        section.right_margin = Inches(settings.get('margin_right', 1.0))

    # --- Default font + spacing on ALL body styles ---
    font_name = settings.get('font', 'Times New Roman')
    font_size = settings.get('font_size', 12)
    ls = settings.get('line_spacing', 2.0)
    space_after = settings.get('space_after', 0)
    space_before = settings.get('space_before', 0)
    fli = settings.get('first_line_indent', 0)

    for sname in body_styles:
        for s in doc.styles:
            if s.name == sname:
                s.font.name = font_name
                s.font.size = Pt(font_size)
                s.font.color.rgb = RGBColor(0, 0, 0)  # Force black, not theme color
                # East-Asian font binding
                rpr = s.element.get_or_add_rPr()
                rfonts = rpr.find(qn('w:rFonts'))
                if rfonts is None:
                    rfonts = OxmlElement('w:rFonts')
                    rpr.insert(0, rfonts)
                rfonts.set(qn('w:eastAsia'), font_name)
                # Line spacing
                pf = s.paragraph_format
                if ls == 2.0:
                    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                elif ls == 1.5:
                    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                elif ls == 1.0:
                    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                else:
                    pf.line_spacing = ls
                pf.space_after = Pt(space_after)
                pf.space_before = Pt(space_before)
                if fli:
                    pf.first_line_indent = Inches(fli)
                break

    # --- Body text alignment ---
    justify = settings.get('justify_body', False)
    if justify:
        for para in doc.paragraphs:
            if para.style.name in ('Body Text', 'First Paragraph', 'Normal') and para.text.strip():
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Widow/orphan control ---
    if settings.get('widow_orphan', True):
        for para in doc.paragraphs:
            pPr = para._p.get_or_add_pPr()
            if pPr.find(qn('w:widowControl')) is None:
                pPr.append(OxmlElement('w:widowControl'))

    # --- Heading font ---
    heading_size = settings.get('heading_size')  # None = auto from body size
    heading_bold = settings.get('heading_bold', True)
    for level in range(1, 7):
        sname = f'Heading {level}'
        for s in doc.styles:
            if s.name == sname:
                s.font.name = font_name
                if heading_size:
                    hs = heading_size
                elif level == 1:
                    hs = int(font_size * 1.3)
                elif level == 2:
                    hs = int(font_size * 1.15)
                else:
                    hs = int(font_size * 1.1)
                s.font.size = Pt(hs)
                s.font.bold = heading_bold
                s.font.color.rgb = RGBColor(0, 0, 0)  # Force black, not theme color
                break

    # --- Line numbering ---
    if settings.get('line_numbers', False):
        restart = settings.get('line_numbers_restart', 'page')
        _add_line_numbering(doc, restart)

    # --- Page numbers ---
    if settings.get('page_numbers', False):
        pos = settings.get('page_numbers_position', 'footer')
        align = settings.get('page_numbers_align', 'center')
        _add_page_numbers(doc, pos, align, font_name, font_size)

    # --- Custom header text ---
    header_text = settings.get('header_text')
    if header_text:
        _add_header_text(doc, header_text, font_name, font_size)

    # --- Caption styling ---
    cap_size = settings.get('caption_font_size', 10)
    cap_align = settings.get('caption_align', 'center')
    _format_captions(doc, cap_size, cap_align, font_name, caption_styles)

    # --- References hanging indent ---
    ref_indent = settings.get('references_hanging_indent', 0.5)
    if ref_indent:
        _format_references(doc, ref_indent, font_name)

    # --- Ensure References heading exists (pandoc sometimes omits it) ---
    _ensure_references_heading(doc)

    # --- Table font sizing ---
    table_font_size = settings.get('table_font_size', 9)
    _format_tables(doc, table_font_size, font_name)

    # --- Table borders (booktabs / none / grid) ---
    table_borders = settings.get('table_borders', 'grid')
    _format_table_borders(doc, table_borders)

    # --- Bold header row on data tables (booktabs visual fix) ---
    _format_table_headers(doc)

    # --- Title page: different first page ---
    if settings.get('title_page', False):
        for section in doc.sections:
            section.different_first_page_header_footer = True

    doc.save(output_path)


def _add_line_numbering(doc, restart='page', count_by=1):
    """Add line numbering to all sections via XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in doc.sections:
        sectPr = section._sectPr
        # Remove existing lnNumType
        for elem in sectPr.findall(qn('w:lnNumType')):
            sectPr.remove(elem)
        # Create new
        ln = OxmlElement('w:lnNumType')
        ln.set(qn('w:countBy'), str(count_by))
        ln.set(qn('w:restart'), restart)
        # Insert after pgMar (correct schema order)
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is not None:
            pgMar.addnext(ln)
        else:
            sectPr.append(ln)


def _add_page_numbers(doc, position='footer', align='center',
                      font_name='Times New Roman', font_size=12):
    """Add page number field to every section's footer or header."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }
    alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    for section in doc.sections:
        target = section.footer if position == 'footer' else section.header
        target.is_linked_to_previous = False
        # Use existing first paragraph or create one
        p = target.paragraphs[0] if target.paragraphs else target.add_paragraph()
        p.alignment = alignment
        # Clear any existing runs
        for run in p.runs:
            run.clear()
        # Build PAGE field
        run = p.add_run()
        run.font.name = font_name
        run.font.size = Pt(font_size)
        fldBegin = OxmlElement('w:fldChar')
        fldBegin.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        fldEnd = OxmlElement('w:fldChar')
        fldEnd.set(qn('w:fldCharType'), 'end')
        run._r.append(fldBegin)
        run._r.append(instr)
        run._r.append(fldEnd)


def _add_header_text(doc, text, font_name='Times New Roman', font_size=10):
    """Add custom running header text."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # Clear existing
        for run in p.runs:
            run.clear()
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.italic = True


def _format_captions(doc, font_size=10, align='center', font_name='Times New Roman',
                     caption_style_names=None):
    """Style figure and table captions."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if caption_style_names is None:
        caption_style_names = ['Caption', 'Image Caption', 'Table Caption']

    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }
    alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    # Update caption styles
    for s in doc.styles:
        if s.name in caption_style_names:
            s.font.size = Pt(font_size)
            s.font.name = font_name
            s.font.italic = False

    # Also walk paragraphs for inline-styled captions
    for para in doc.paragraphs:
        text = para.text.strip()
        if para.style.name in caption_style_names:
            for run in para.runs:
                run.font.size = Pt(font_size)
                run.font.name = font_name
            para.alignment = alignment
        elif text.startswith(('Figure ', 'Table ', 'Fig. ', 'Tab. ')):
            import re
            if re.match(r'^(Figure|Table|Fig\.|Tab\.)\s+\d+', text):
                for run in para.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = font_name
                para.alignment = alignment


def _format_references(doc, hanging_indent=0.5, font_name='Times New Roman'):
    """Apply hanging indent to References/Bibliography entries."""
    from docx.shared import Inches, Pt

    # Strategy 1: Update 'Bibliography' style directly (pandoc uses this)
    bib_style_found = False
    for s in doc.styles:
        if s.name in ('Bibliography', 'References'):
            pf = s.paragraph_format
            pf.left_indent = Inches(hanging_indent)
            pf.first_line_indent = Inches(-hanging_indent)
            s.font.name = font_name
            bib_style_found = True
            break

    # Strategy 2: Walk paragraphs looking for References heading
    # (for documents that don't use the Bibliography style)
    if not bib_style_found:
        in_refs = False
        for para in doc.paragraphs:
            text = para.text.strip().lower()
            style_name = para.style.name

            if text in ('references', 'bibliography', 'literature cited') or \
               (text.startswith('references') and len(text) < 40):
                in_refs = True
                continue

            if in_refs:
                if style_name.startswith('Heading') or \
                   text.startswith(('appendix', 'acknowledg')):
                    in_refs = False
                    continue
                if para.text.strip():
                    pf = para.paragraph_format
                    pf.left_indent = Inches(hanging_indent)
                    pf.first_line_indent = Inches(-hanging_indent)
                    for run in para.runs:
                        run.font.name = font_name


def _format_tables(doc, font_size=9, font_name='Times New Roman'):
    """Set font size on all table cells."""
    from docx.shared import Pt

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(font_size)
                        run.font.name = font_name


def _format_table_borders(doc, style='booktabs'):
    """Strip cell borders and apply clean booktabs, grid, or borderless style.

    Only touches real data tables (≥2 rows with text content).
    Single-row tables (figure/image wrappers) are skipped.

    Args:
        style: 'booktabs' — top/bottom table rules + thin rule below header.
               'none' — no borders at all.
               'grid' — full grid with thin (0.5 pt) borders on all sides.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for table in doc.tables:
        # --- Skip non-data tables (image wrappers, empty tables) ---
        if len(table.rows) < 2:
            continue
        has_text = any(cell.text.strip()
                       for row in table.rows
                       for cell in row.cells)
        if not has_text:
            continue

        # 1. Remove existing cell-level tcBorders (use find, NOT get_or_add)
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)

        # 2. Remove existing table-level borders
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            tblPr.remove(tblBorders)

        if style == 'none':
            continue  # Done — borders fully removed

        if style == 'grid':
            # Grid: all table borders (0.5 pt) — standard Elsevier table style
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')     # 0.5 pt
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            continue  # Done — grid borders added, skip booktabs header step

        # Booktabs: top + bottom table borders (1.5 pt)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name, sz in [('top', 12), ('bottom', 12)]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), str(sz))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        # Thin bottom border on header row (1 pt) — only for booktabs
        if table.rows and table.rows[0].cells:
            header_row = table.rows[0]
            for cell in header_row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    cell._tc.insert(0, tcPr)
                tcBorders = OxmlElement('w:tcBorders')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '8')    # 1 pt
                bottom.set(qn('w:space'), '0')
                bottom.set(qn('w:color'), '000000')
                tcBorders.append(bottom)
                tcPr.append(tcBorders)


def _format_table_headers(doc):
    """Bold the first row of every data table (visual header in booktabs style)."""
    for table in doc.tables:
        if len(table.rows) < 2:
            continue
        if not any(cell.text.strip()
                   for row in table.rows for cell in row.cells):
            continue
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True


def _ensure_references_heading(doc):
    """Add a 'References' heading if no heading exists before Bibliography entries."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Find the first Bibliography paragraph
    first_bib_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == 'Bibliography':
            first_bib_idx = i
            break
    if first_bib_idx is None:
        return  # No bibliography, nothing to do

    # Check if there's already a heading right before the first bibliography entry
    prev = doc.paragraphs[first_bib_idx - 1] if first_bib_idx > 0 else None
    if prev and prev.style.name.startswith('Heading'):
        return  # Already has a heading

    # Check if any heading with "Reference" text exists nearby
    for i in range(max(0, first_bib_idx - 3), first_bib_idx):
        p = doc.paragraphs[i]
        if p.style.name.startswith('Heading') and \
           p.text.strip().lower() in ('references', 'bibliography'):
            return  # Heading exists but might not be immediately before bib

    # Need to insert a References heading before the first bibliography entry
    bib_para = doc.paragraphs[first_bib_idx]
    # Get the element to insert before
    bib_p_elem = bib_para._p

    # Create a new paragraph with Heading 1 style
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Heading1')
    pPr.append(pStyle)
    new_p.append(pPr)

    # Add the text "References"
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rPr2 = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = 'References'
    r.append(t)
    new_p.append(r)

    # Insert before the first bibliography paragraph
    parent = bib_p_elem.getparent()
    if parent is not None:
        parent.insert(list(parent).index(bib_p_elem), new_p)


def main():
    parser = argparse.ArgumentParser(
        description='Prepare LaTeX for pandoc → Word AND/OR format .docx to journal specs',
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    # === Input/Output (shared) ===
    parser.add_argument('input', help='Input .tex (LaTeX prep) or .docx (format mode)')
    parser.add_argument('-o', '--output', help='Output file path')

    # === Mode selection ===
    parser.add_argument('--format-docx', action='store_true',
                        help='Format an existing .docx instead of prepping LaTeX '
                             '(auto-detected if input ends in .docx)')

    # === LaTeX prep options ===
    tex_group = parser.add_argument_group('LaTeX Prep')
    tex_group.add_argument('--elsarticle', action='store_true',
                           help='Transform elsarticle preamble → article class')
    tex_group.add_argument('--dry-run', action='store_true',
                           help='Show label mapping without modifying files')
    tex_group.add_argument('--no-refs', action='store_true',
                           help='Skip \\ref{} replacement')
    tex_group.add_argument('--no-captions', action='store_true',
                           help='Skip caption prefixing')
    tex_group.add_argument('--no-tables', action='store_true',
                           help='Skip table fixes')

    # === DOCX: Journal presets ===
    preset_group = parser.add_argument_group('DOCX Journal Presets')
    preset_group.add_argument('--journal',
                              choices=list(JOURNAL_PRESETS.keys()),
                              metavar='NAME',
                              help='Journal preset: '
                                   + ', '.join(
                                       f'{k} ({v["description"]})'
                                       for k, v in JOURNAL_PRESETS.items()))

    # === DOCX: Page setup ===
    page_group = parser.add_argument_group('DOCX Page Setup')
    page_group.add_argument('--page-size', choices=list(PAGE_SIZES.keys()),
                            metavar='SIZE',
                            help='Page size: ' + ', '.join(PAGE_SIZES.keys()))
    page_group.add_argument('--margins', type=float, metavar='IN',
                            help='All four margins (inches)')
    page_group.add_argument('--margin-top', type=float, metavar='IN')
    page_group.add_argument('--margin-bottom', type=float, metavar='IN')
    page_group.add_argument('--margin-left', type=float, metavar='IN')
    page_group.add_argument('--margin-right', type=float, metavar='IN')

    # === DOCX: Typography ===
    typo_group = parser.add_argument_group('DOCX Typography')
    typo_group.add_argument('--font', metavar='NAME',
                            help='Body font (e.g., "Times New Roman")')
    typo_group.add_argument('--font-size', type=float, metavar='PT',
                            help='Body font size (default: 12)')
    typo_group.add_argument('--line-spacing', type=float, metavar='N',
                            help='Line spacing: 1.0, 1.5, 2.0 (default: 2.0)')
    typo_group.add_argument('--space-after', type=float, metavar='PT',
                            help='Space after paragraphs (default: 0)')
    typo_group.add_argument('--space-before', type=float, metavar='PT',
                            help='Space before paragraphs (default: 0)')
    typo_group.add_argument('--first-line-indent', type=float, metavar='IN',
                            help='First-line indent (default: 0)')
    typo_group.add_argument('--justify', action='store_true', default=None,
                            help='Justify body text')
    typo_group.add_argument('--no-justify', action='store_false', dest='justify',
                            help='Left-align body text (overrides preset)')

    # === DOCX: Numbering ===
    num_group = parser.add_argument_group('DOCX Numbering')
    num_group.add_argument('--line-numbers', action='store_true', default=None,
                           help='Enable line numbering')
    num_group.add_argument('--no-line-numbers', action='store_false', dest='line_numbers',
                           help='Disable line numbering (overrides preset)')
    num_group.add_argument('--line-numbers-restart',
                           choices=['page', 'continuous'], default='page',
                           help='Line number restart: per page or continuous')
    num_group.add_argument('--page-numbers', action='store_true', default=None,
                           help='Enable page numbers')
    num_group.add_argument('--no-page-numbers', action='store_false', dest='page_numbers',
                           help='Disable page numbers (overrides preset)')
    num_group.add_argument('--page-numbers-position',
                           choices=['header', 'footer'], default='footer',
                           help='Page number position (default: footer)')
    num_group.add_argument('--page-numbers-align',
                           choices=['center', 'left', 'right'], default='center',
                           help='Page number alignment (default: center)')

    # === DOCX: Element styling ===
    style_group = parser.add_argument_group('DOCX Element Styling')
    style_group.add_argument('--caption-size', type=int, metavar='PT',
                             help='Caption font size (default: 10)')
    style_group.add_argument('--caption-align',
                             choices=['center', 'left', 'right'],
                             help='Caption alignment (default: center)')
    style_group.add_argument('--ref-indent', type=float, metavar='IN',
                             help='References hanging indent (default: 0.5)')
    style_group.add_argument('--table-size', type=int, metavar='PT',
                              help='Table cell font size (default: 9)')
    style_group.add_argument('--table-borders',
                              choices=['booktabs', 'none', 'grid'],
                              help='Table border style: booktabs (top/bottom rules), none, grid')
    style_group.add_argument('--heading-bold', action='store_true', default=None,
                              help='Bold section headings')
    style_group.add_argument('--no-heading-bold', action='store_false', dest='heading_bold',
                              help='Regular (non-bold) headings')
    style_group.add_argument('--heading-size', type=int, metavar='PT',
                              help='Heading font size (default: auto — body × 1.3/1.15)')

    # === DOCX: Headers/footers ===
    hf_group = parser.add_argument_group('DOCX Headers/Footers')
    hf_group.add_argument('--header-text', metavar='TEXT',
                          help='Custom running header text')
    hf_group.add_argument('--title-page', action='store_true', default=None,
                          help='Different first page (title page)')
    hf_group.add_argument('--no-widow-orphan', action='store_false',
                          default=None, dest='widow_orphan',
                          help='Disable widow/orphan control')

    args = parser.parse_args()

    # =========================================================================
    # DOCX FORMAT MODE
    # =========================================================================
    if args.format_docx or args.input.lower().endswith('.docx'):
        # Build settings dict — start with preset, then override with explicit args
        settings = {}
        if args.journal:
            settings = JOURNAL_PRESETS[args.journal].copy()
            settings.pop('description', None)
            print(f"Using preset: {args.journal} — {JOURNAL_PRESETS[args.journal]['description']}")

        # Override individual settings (only non-None args)
        overrides = {
            'page_size': args.page_size,
            'font': args.font,
            'font_size': args.font_size,
            'line_spacing': args.line_spacing,
            'margin_top': args.margin_top or (args.margins if args.margins else None),
            'margin_bottom': args.margin_bottom or (args.margins if args.margins else None),
            'margin_left': args.margin_left or (args.margins if args.margins else None),
            'margin_right': args.margin_right or (args.margins if args.margins else None),
            'justify_body': args.justify,
            'first_line_indent': args.first_line_indent,
            'space_after': args.space_after,
            'space_before': args.space_before,
            'line_numbers': args.line_numbers,
            'line_numbers_restart': args.line_numbers_restart if args.line_numbers_restart != 'page' else None,
            'page_numbers': args.page_numbers,
            'page_numbers_position': args.page_numbers_position if args.page_numbers_position != 'footer' else None,
            'page_numbers_align': args.page_numbers_align if args.page_numbers_align != 'center' else None,
            'caption_font_size': args.caption_size,
            'caption_align': args.caption_align,
            'references_hanging_indent': args.ref_indent,
            'table_font_size': args.table_size,
            'table_borders': args.table_borders,
            'heading_bold': args.heading_bold,
            'heading_size': args.heading_size,
            'header_text': args.header_text,
            'title_page': args.title_page,
            'widow_orphan': args.widow_orphan,
        }
        for k, v in overrides.items():
            if v is not None:
                settings[k] = v

        # Default output path
        if not args.output:
            base = args.input.rsplit('.', 1)[0]
            args.output = f"{base}_formatted.docx"

        format_docx(args.input, args.output, settings)

        # Print summary
        applied = []
        checks = [
            ('page_size', lambda s: f"page={s.get('page_size', 'letter')}"),
            ('font', lambda s: f"font={s.get('font', 'Times New Roman')} {s.get('font_size', 12)}pt"),
            ('spacing', lambda s: f"line-spacing={s.get('line_spacing', 2.0)}"),
            ('margins', lambda s: f"margins={s.get('margin_top', 1.0)}/{s.get('margin_bottom', 1.0)}/{s.get('margin_left', 1.0)}/{s.get('margin_right', 1.0)}\""),
            ('line_nums', lambda s: f"line-numbers={'ON' if s.get('line_numbers') else 'OFF'}"),
            ('page_nums', lambda s: f"page-numbers={'ON' if s.get('page_numbers') else 'OFF'}"),
            ('justify', lambda s: f"justify={'ON' if s.get('justify_body') else 'OFF'}"),
            ('captions', lambda s: f"captions={s.get('caption_font_size', 10)}pt"),
            ('tables', lambda s: f"tables={s.get('table_font_size', 9)}pt"),
            ('tborders', lambda s: f"table-borders={s.get('table_borders', 'grid')}"),
            ('headings', lambda s: f"headings={'bold' if s.get('heading_bold', True) else 'regular'}"),
            ('refs', lambda s: f"ref-indent={s.get('references_hanging_indent', 0.5)}\""),
        ]
        for _, fn in checks:
            applied.append(fn(settings))

        print(f"\n✓ Formatted: {args.output}")
        for item in applied:
            print(f"  {item}")
        return

    # =========================================================================
    # LATEX PREP MODE
    # =========================================================================
    with open(args.input, 'r') as f:
        content = f.read()

    base_dir = os.path.dirname(os.path.abspath(args.input))

    # Transform preamble first (before resolving inputs)
    if args.elsarticle:
        content = transform_elsarticle_preamble(content)

    # Build search paths and resolve \input{} directives early
    # so all fixes operate on the fully expanded document
    input_paths = build_input_paths(content, base_dir)
    content = resolve_inputs(content, base_dir, input_paths)

    # Detect labels from the full expanded document
    labels = detect_labels(content)
    fig_map, tab_map = build_label_map(labels)

    if args.dry_run:
        print(extract_label_map_text(labels, fig_map, tab_map))
        return

    # Apply fixes on the fully expanded content
    if not args.no_refs:
        content = fix_refs(content, fig_map, tab_map)

    if not args.no_captions:
        content = fix_captions(content, fig_map, tab_map)

    if not args.no_tables:
        content = fix_longtables(content)
        content = fix_resizebox(content)
        content = fix_table_columns(content)
        content = fix_table_math(content)
        content = fix_addlinespace(content)

    # Always fix double prefixes (result of ref replacement)
    content = fix_double_prefixes(content)

    # Determine output path
    if args.output:
        out_path = args.output
    else:
        base = args.input.rsplit('.', 1)[0]
        out_path = f"{base}_pandoc.tex"

    with open(out_path, 'w') as f:
        f.write(content)

    print(f"Output: {out_path}")
    print(extract_label_map_text(labels, fig_map, tab_map))
    print(f"\n{'✓' if args.elsarticle else '○'} Preamble: {'transformed elsarticle → article' if args.elsarticle else 'kept as-is'}")
    print(f"{'✓' if not args.no_refs else '○'} Refs: replaced \\ref with literal numbers")
    print(f"{'✓' if not args.no_captions else '○'} Captions: added Figure N:/Table N: prefixes")
    print(f"{'✓' if not args.no_tables else '○'} Tables: longtable→tabular + column specs + math + addlinespace")
    print(f"\nNext: pandoc {out_path} -o output.docx --citeproc --bibliography=refs.bib --csl=apa.csl")


if __name__ == '__main__':
    main()
